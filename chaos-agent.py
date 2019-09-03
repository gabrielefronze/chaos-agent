from docx import Document
from random import shuffle

from kivy.config import Config
Config.set('graphics', 'resizable', False)
Config.set('graphics', 'resizable', '0')  # 0 being off 1 being on as in true/false
Config.set('graphics', 'width', '500')
Config.set('graphics', 'height', '500')
Config.write()

from kivy.app import App
from kivy.uix.button import Button
from kivy.uix.floatlayout import FloatLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.slider import Slider
from kivy.core.window import Window
from kivy.properties import NumericProperty
from kivy.uix.popup import Popup
from kivy.uix.filechooser import FileChooserIconView

import os
import shutil


class FileSet:
    def __init__(self, filename="test.docx", size=1):

        self.original_file_name = filename
        self.original_file = Document(filename)

        path = filename.rsplit('/', 1)[0]
        filename = filename.rsplit('/', 1)[1]

        print(filename)
        print(path)
        print(self.original_file_name)

        prefix = os.path.join(path, 'shuffled')
        prefix = os.path.join(prefix, '')

        print(prefix)

        try:
            print('removing')
            shutil.rmtree(prefix)
        except:
            pass

        try:
            print('recreating')
            os.mkdir(prefix)
        except:
            pass

        self.shuffled_files_names = []
        for i in range(0, size):
            shuffled_filename = os.path.join(prefix, filename.replace(".docx", "_{}.docx".format(i+1)))
            print(shuffled_filename)
            self.shuffled_files_names.append(shuffled_filename)

        self.shuffled_files = []
        for fn in self.shuffled_files_names:
            self.original_file.save(fn)
            self.shuffled_files.append(Document(fn))

    def shuffle(self):

        for index in range(0, len(self.shuffled_files)):
            texts = {}
            was_list = False

            for i, p in enumerate(self.shuffled_files[index].paragraphs):
                if "List" in p.style.name:
                    texts[i] = p.text
                    was_list = True
                    # print("{}: That's a list paragraph!".format(i))
                elif was_list or i == len(self.shuffled_files[index].paragraphs)-1:

                    print("mixin'")

                    old_indexes = list(texts.keys())
                    new_indexes = old_indexes.copy()
                    new_indexes.sort()
                    shuffle(old_indexes)
                    shuffle(new_indexes)

                    index_conversion = dict(zip(old_indexes, new_indexes))
                    print(index_conversion)

                    k = 0
                    for j, text in texts.items():
                        self.shuffled_files[index].paragraphs[j].text = texts[index_conversion[j]]

                    texts.clear()
                    was_list = False

    def save(self):
        for i, sf in enumerate(self.shuffled_files):
            sf.save(self.shuffled_files_names[i])


class Container(FloatLayout):
    slider_val = NumericProperty(25)

    # layout
    def __init__(self, *args, **kwargs):
        super(Container, self).__init__(*args, **kwargs)
        Window.clearcolor = (1, 1, 1, 1)
        Window.color = (0, 0, 0, 1)
        Window.size=(800, 400)
        font_color = (0, 0, 0, 1)
        self.size=(300, 50)

        self.txt1 = TextInput(text='',
                              multiline=False,
                              size_hint=(.5, .07),
                              pos_hint={'x': .3, 'y': .7})

        self.openpopupbtn = Button(text='Select',
                                   size_hint=(.1, .07),
                                   pos_hint={'x': .8, 'y': .7})

        self.filechs = FileChooserIconView(filters=['*.docx'],
                                           path='./',
                                           size_hint=(.8, .7),
                                           pos_hint={'x': .1, 'y': .3})

        self.popupcontent = FloatLayout()
        self.popupcontent.add_widget(self.filechs)

        self.popup = Popup(title='Select file',
                           content=self.popupcontent,
                           size_hint=(.8, .8))

        self.openpopupbtn.bind(on_press=self.popup.open)

        self.closepopupbtn = Button(text='Select',
                                    size_hint=(.8, .1),
                                    pos_hint={'x': .1, 'y': .1})
        self.closepopupbtn.bind(on_press=self.selectfile)

        self.popupcontent.add_widget(self.closepopupbtn)

        self.add_widget(Label(text="Original file name:",
                                size_hint=(.4, .05),
                                pos_hint={'x': .0, 'y': .7},
                                color=font_color))
        self.add_widget(self.txt1)
        self.add_widget(self.openpopupbtn)

        self.slider = Slider(min=1,
                             max=40,
                             value=25,
                             step=1,
                             size_hint=(.5, .05),
                             pos_hint={'x': .3, 'y': .5})
        self.slider.bind(value=self.on_value)

        self.lblslider = Label(text="25",
                               size_hint=(.3, .05),
                               pos_hint={'x':.7, 'y':.5},
                               color=font_color)

        self.add_widget(Label(text="Number of output files:",
                                size_hint=(.4, .05),
                                pos_hint={'x': .0, 'y': .5},
                                color=font_color))
        self.add_widget(self.slider)
        self.add_widget(self.lblslider)

        btn1 = Button(text="Create!",
                      size_hint=(1, .15),
                      pos_hint={'x':.0, 'y':.2})
        btn1.bind(on_press=self.buttonClicked)

        self.lbl1 = Label(text="",
                          size_hint=(1, .05),
                          pos_hint={'x': .0, 'y': .1},
                          color=font_color)

        self.add_widget(self.lbl1)
        # layout.add_widget(self.txt2)

        self.add_widget(btn1)

    # button click function
    def buttonClicked(self, btn):
        filename = self.txt1.text
        size = int(self.slider.value)

        self.lbl1.text = "Creating {} files from {} original file.".format(size, filename)

        shuffler = FileSet(filename=filename, size=size)
        shuffler.shuffle()
        shuffler.save()

    def on_value(self, instance, value):
        print(value)
        self.lblslider.text = str(value)

    def selectfile(self, btn):
        self.txt1.text = os.path.join(self.filechs.path, self.filechs.selection[0])
        self.popup.dismiss()


class MainApp(App):
    def build(self):
        self.icon = './icon.ico'
        self.title = 'Chaos agent - list shuffler'
        return Container()


if __name__ == '__main__':
    MainApp().run()
